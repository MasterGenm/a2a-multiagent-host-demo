# -*- coding: utf-8 -*-
from __future__ import annotations
import os

from docx import Document as Docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from ..model import DocumentModel, Section, TableData
from .base import BaseWriter, ensure_parent

_CJK_FONT_CANDIDATES = [
    ("Microsoft YaHei", r"C:\Windows\Fonts\msyh.ttc"),
    ("SimSun", r"C:\Windows\Fonts\simsun.ttc"),
]


def _set_style_font(style, font_name: str) -> None:
    try:
        style.font.name = font_name
    except Exception:
        return

    try:
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rfonts.set(qn("w:eastAsia"), font_name)
        rfonts.set(qn("w:cs"), font_name)
    except Exception:
        pass


def _pick_cjk_font_name() -> str:
    for name, font_path in _CJK_FONT_CANDIDATES:
        try:
            if os.path.exists(font_path):
                return name
        except Exception:
            pass
    return _CJK_FONT_CANDIDATES[0][0]


class DocxWriter(BaseWriter):
    def write(self, doc: DocumentModel, output_path: str) -> str:
        ensure_parent(output_path)
        d = Docx()

        # 为中文设置一个稳定的默认字体（避免某些环境里中文显示异常/字体回退不一致）
        # Word 会自动回退到系统可用字体；这里尽量选一个本机确定存在的字体名。
        cjk_font = _pick_cjk_font_name()
        for style_name in (
            "Normal",
            "Title",
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "List Bullet",
            "List Number",
        ):
            try:
                _set_style_font(d.styles[style_name], cjk_font)
            except Exception:
                pass

        # 封面/标题
        title_p = d.add_paragraph()
        run = title_p.add_run(doc.meta.title)
        run.bold = True; run.font.size = Pt(20)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if doc.meta.subtitle:
            p = d.add_paragraph(doc.meta.subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_line = f"{doc.meta.author or ''}    {doc.meta.date or ''}".strip()
        if meta_line:
            p = d.add_paragraph(meta_line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        d.add_page_break()

        # 正文
        for s in doc.sections:
            d.add_heading(s.title, level=1)
            for para in s.paragraphs:
                d.add_paragraph(para)
            if s.bullets:
                for b in s.bullets:
                    d.add_paragraph(b, style="List Bullet")
            for t in s.tables:
                self._add_table(d, t)
            # 图片可选
            for f in s.figures:
                try:
                    d.add_picture(f.path, width=Inches(6))
                    if f.caption:
                        cap = d.add_paragraph(f.caption)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass

        if doc.references:
            d.add_heading("参考文献", level=1)
            for r in doc.references:
                d.add_paragraph(r, style="List Number")

        d.save(output_path)
        return output_path

    @staticmethod
    def _add_table(d: Docx, t: TableData):
        rows = len(t.rows) + (1 if t.headers else 0)
        cols = len(t.headers) if t.headers else max((len(r) for r in t.rows), default=1)
        table = d.add_table(rows=rows, cols=cols)
        if t.headers:
            hdr = table.rows[0].cells
            for j, h in enumerate(t.headers):
                hdr[j].text = str(h)
        for i, row in enumerate(t.rows):
            cells = table.rows[i + (1 if t.headers else 0)].cells
            for j, val in enumerate(row):
                cells[j].text = str(val)
